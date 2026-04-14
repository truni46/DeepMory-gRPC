# server/modules/knowledge/service.py
import asyncio
import hashlib
import os
import uuid
from pathlib import Path
from typing import Dict, List, Optional

from config.logger import logger
from modules.knowledge.repository import documentRepository
from modules.rag.ragService import ragService
from modules.llm.llmProvider import llmProvider

UPLOAD_DIR = Path(__file__).parent.parent.parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_TYPES = {".pdf", ".txt", ".md", ".docx", ".doc", ".xlsx", ".xls"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
SUMMARY_MAX_CHARS = 4000


def _computeHash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _extractPageCount(filePath: str) -> int:
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            return len(reader.pages)
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(filePath)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            return max(len(paragraphs) // 25, 1)
    except Exception as e:
        logger.warning(f"_extractPageCount failed for {filePath}: {e}")
    return 0


def _readTextContent(filePath: str, maxChars: int = SUMMARY_MAX_CHARS) -> str:
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(filePath)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            with open(filePath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        return text[:maxChars]
    except Exception as e:
        logger.warning(f"_readTextContent failed for {filePath}: {e}")
        return ""


class DocumentService:

    async def uploadDocuments(
        self,
        userId: str,
        files: List,
        scope: str = "personal",
        ownerId: str = None,
        ownerType: str = "user",
    ) -> List[Dict]:
        results = []
        for fileObj in files:
            try:
                doc = await self._uploadOne(
                    userId, fileObj, scope, ownerId or userId, ownerType
                )
                results.append(doc)
            except Exception as e:
                logger.error(f"uploadDocuments failed for {fileObj.filename}: {e}")
                results.append({"error": str(e), "filename": fileObj.filename})
        return results

    async def _uploadOne(
        self,
        userId: str,
        fileObj,
        scope: str,
        ownerId: str,
        ownerType: str,
    ) -> Dict:
        filename = fileObj.filename
        fileExt = os.path.splitext(filename)[1].lower()
        if fileExt not in ALLOWED_TYPES:
            raise ValueError(f"File type '{fileExt}' not allowed")

        content = await fileObj.read()
        if len(content) > MAX_FILE_SIZE:
            raise ValueError("File exceeds 50 MB limit")

        contentHash = _computeHash(content)
        storedFilename = f"{uuid.uuid4().hex}_{filename}"
        filePath = UPLOAD_DIR / storedFilename

        with open(filePath, "wb") as f:
            f.write(content)

        record = await documentRepository.create(
            userId=userId,
            filename=filename,
            storedFilename=storedFilename,
            filePath=str(filePath),
            fileType=fileExt.lstrip("."),
            fileSize=len(content),
            contentHash=contentHash,
            scope=scope,
            ownerId=ownerId,
            ownerType=ownerType,
        )

        asyncio.create_task(
            self._processDocument(record["id"], str(filePath), ownerId, userId)
        )
        return record

    async def _processDocument(
        self, documentId: str, filePath: str, ownerId: str, userId: str
    ) -> None:
        try:
            await documentRepository.updateEmbedding(documentId, "processing")
            await ragService.index(filePath, ownerId, documentId, userId)
            pageCount = _extractPageCount(filePath)
            await documentRepository.updateEmbedding(
                documentId, "completed", chunkCount=1, pageCount=pageCount
            )
            logger.info(f"_processDocument completed for {documentId}")
            asyncio.create_task(self._generateSummary(documentId, filePath))
        except Exception as e:
            logger.error(f"_processDocument failed for {documentId}: {e}")
            await documentRepository.updateEmbedding(
                documentId, "failed", errorMsg=str(e)
            )

    async def _generateSummary(self, documentId: str, filePath: str) -> None:
        try:
            await documentRepository.updateSummary(documentId, "processing")
            text = _readTextContent(filePath)
            if not text.strip():
                await documentRepository.updateSummary(documentId, "failed")
                return
            messages = [
                {
                    "role": "system",
                    "content": "You are a helpful assistant that summarizes documents concisely.",
                },
                {
                    "role": "user",
                    "content": f"Summarize the following document in 3-5 sentences:\n\n{text}",
                },
            ]
            response = await llmProvider.generateResponse(messages, stream=False)
            await documentRepository.updateSummary(
                documentId, "completed", summary=response
            )
            logger.info(f"_generateSummary completed for {documentId}")
        except Exception as e:
            logger.error(f"_generateSummary failed for {documentId}: {e}")
            await documentRepository.updateSummary(documentId, "failed")

    async def getDocuments(
        self, userId: str, scope: Optional[str] = None
    ) -> List[Dict]:
        return await documentRepository.getByUser(userId, scope)

    async def getDocument(self, documentId: str, userId: str) -> Optional[Dict]:
        doc = await documentRepository.getById(documentId)
        if not doc or str(doc.get("userId")) != str(userId):
            return None
        return doc

    async def getDocumentContext(self, documentIds: List[str], userId: str) -> str:
        parts = []
        for docId in documentIds:
            try:
                doc = await self.getDocument(docId, userId)
                if not doc:
                    logger.warning(f"getDocumentContext: document {docId} not found or unauthorized for user {userId}")
                    continue
                filePath = doc.get("filePath", "")
                if not filePath:
                    logger.warning(f"getDocumentContext: no filePath for document {docId}")
                    continue
                text = _readTextContent(filePath, maxChars=8000)
                if text.strip():
                    parts.append(f"--- Document: {doc.get('filename', docId)} ---\n{text}")
            except Exception as e:
                logger.error(f"getDocumentContext failed for document {docId}: {e}")
        return "\n\n".join(parts)

    async def deleteDocument(self, userId: str, documentId: str) -> bool:
        result = await documentRepository.delete(documentId, userId)
        if result is None:
            return False
        filePath, ownerId = result
        try:
            if filePath and os.path.exists(filePath):
                os.remove(filePath)
        except Exception as e:
            logger.error(f"deleteDocument file removal failed for {documentId}: {e}")
        try:
            await ragService.deleteDocumentChunks(ownerId, documentId)
        except Exception as e:
            logger.error(f"deleteDocument RAG cleanup failed for {documentId}: {e}")
        return True


documentService = DocumentService()
