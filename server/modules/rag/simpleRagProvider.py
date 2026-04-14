from __future__ import annotations

import os
import uuid
from typing import Dict, List, Optional

from qdrant_client import AsyncQdrantClient
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    FilterSelector,
    MatchValue,
    PointIdsList,
    PointStruct,
    VectorParams,
)

from config.logger import logger
from modules.llm.embeddingProvider import embeddingService
from modules.rag.repository import Document, SearchResult


def _readFile(filePath: str) -> str:
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(filePath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            with open(filePath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        logger.error(f"_readFile failed for '{filePath}': {e}")
        return ""


def _chunkText(text: str, chunkSize: int = 800, overlap: int = 100) -> List[str]:
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        chunks.append(text[start:start + chunkSize])
        start += chunkSize - overlap
    return [c for c in chunks if c.strip()]


class SimpleRagProvider:
    """Direct Qdrant + embedding RAG — no external graph dependencies."""

    def __init__(self):
        self._url = os.getenv("QDRANT_URL", "http://localhost:6333")
        self._apiKey = os.getenv("QDRANT_API_KEY") or None
        self._client: Optional[AsyncQdrantClient] = None

    async def _getClient(self) -> AsyncQdrantClient:
        if self._client is None:
            self._client = AsyncQdrantClient(url=self._url, api_key=self._apiKey)
        return self._client

    def _collectionName(self, namespace: str) -> str:
        return f"simple_rag_{namespace}"

    async def _ensureCollection(self, name: str) -> None:
        client = await self._getClient()
        collections = await client.get_collections()
        existing = {c.name for c in collections.collections}
        if name not in existing:
            await client.create_collection(
                collection_name=name,
                vectors_config=VectorParams(
                    size=embeddingService.dimension,
                    distance=Distance.COSINE,
                ),
            )
            logger.info(f"SimpleRagProvider: created collection '{name}'")

    async def index(self, filePath: str, projectId: str, documentId: str, userId: str) -> int:
        try:
            content = _readFile(filePath)
            if not content.strip():
                logger.warning(f"SimpleRagProvider: no content extracted from '{filePath}'")
                return 0
            chunks = _chunkText(content)
            if not chunks:
                return 0
            vectors = await embeddingService.embedBatch(chunks)
            collName = self._collectionName(f"project_{projectId}")
            await self._ensureCollection(collName)
            client = await self._getClient()
            points = [
                PointStruct(
                    id=str(uuid.uuid4()),
                    vector=vectors[i],
                    payload={
                        "documentId": documentId,
                        "userId": userId,
                        "chunkIndex": i,
                        "text": chunks[i],
                    },
                )
                for i in range(len(chunks))
            ]
            await client.upsert(collection_name=collName, points=points)
            logger.info(f"SimpleRagProvider: indexed {len(points)} chunks for document {documentId}")
            return len(points)
        except Exception as e:
            logger.error(f"SimpleRagProvider.index failed for document {documentId}: {e}")
            raise

    async def deleteDocumentChunks(self, projectId: str, documentId: str) -> None:
        try:
            client = await self._getClient()
            collName = self._collectionName(f"project_{projectId}")
            await client.delete(
                collection_name=collName,
                points_selector=FilterSelector(
                    filter=Filter(
                        must=[FieldCondition(key="documentId", match=MatchValue(value=documentId))]
                    )
                ),
            )
            logger.info(f"SimpleRagProvider: deleted chunks for document {documentId}")
        except Exception as e:
            logger.error(f"SimpleRagProvider.deleteDocumentChunks failed for document {documentId}: {e}")

    async def searchContext(
        self, query: str, projectId: str, limit: int = 5, **kwargs
    ) -> List[SearchResult]:
        try:
            client = await self._getClient()
            collName = self._collectionName(f"project_{projectId}")
            queryVector = await embeddingService.embed(query)
            results = await client.search(
                collection_name=collName,
                query_vector=queryVector,
                limit=limit,
                with_payload=True,
            )
            return [
                SearchResult(
                    document=Document(
                        id=r.payload.get("documentId", str(uuid.uuid4())),
                        content=r.payload.get("text", ""),
                        metadata={
                            "documentId": r.payload.get("documentId"),
                            "chunkIndex": r.payload.get("chunkIndex"),
                        },
                    ),
                    score=r.score,
                )
                for r in results
                if r.payload.get("text")
            ]
        except Exception as e:
            logger.warning(f"SimpleRagProvider.searchContext failed for project {projectId}: {e}")
            return []

    async def upsertMemoryVector(
        self, userId: str, memoryId: str, content: str, metadata: Optional[Dict] = None
    ) -> None:
        try:
            collName = self._collectionName(f"user_{userId}")
            await self._ensureCollection(collName)
            client = await self._getClient()
            vector = await embeddingService.embed(content)
            await client.upsert(
                collection_name=collName,
                points=[
                    PointStruct(
                        id=memoryId,
                        vector=vector,
                        payload={"text": content, **(metadata or {})},
                    )
                ],
            )
        except Exception as e:
            logger.error(f"SimpleRagProvider.upsertMemoryVector failed for user {userId}: {e}")

    async def searchMemoryVectors(
        self, userId: str, query: str, limit: int = 5
    ) -> List[SearchResult]:
        try:
            client = await self._getClient()
            collName = self._collectionName(f"user_{userId}")
            queryVector = await embeddingService.embed(query)
            results = await client.search(
                collection_name=collName,
                query_vector=queryVector,
                limit=limit,
                with_payload=True,
            )
            return [
                SearchResult(
                    document=Document(
                        id=str(r.id),
                        content=r.payload.get("text", ""),
                        metadata={},
                    ),
                    score=r.score,
                )
                for r in results
                if r.payload.get("text")
            ]
        except Exception as e:
            logger.warning(f"SimpleRagProvider.searchMemoryVectors failed for user {userId}: {e}")
            return []

    async def deleteMemoryVector(self, userId: str, memoryId: str) -> None:
        try:
            client = await self._getClient()
            collName = self._collectionName(f"user_{userId}")
            await client.delete(
                collection_name=collName,
                points_selector=PointIdsList(points=[memoryId]),
            )
        except Exception as e:
            logger.error(f"SimpleRagProvider.deleteMemoryVector failed memoryId={memoryId}: {e}")


simpleRagProvider = SimpleRagProvider()
